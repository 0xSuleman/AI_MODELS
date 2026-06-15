"""Generate the Experiment Runbook DOCX document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import itertools

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
        set_cell_shading(c, '1E3A5F')
        for r in c.paragraphs:
            for run in r.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()
    return t

def add_info_box(doc, label, text, bold_label=True):
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(f'{label}: ')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_run_entry(doc, run_num, total, config_dict, notes=""):
    p = doc.add_paragraph()
    r = p.add_run(f'Run {run_num}/{total}')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    if notes:
        r2 = p.add_run(f'  —  {notes}')
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    code = "ExperimentConfig(\n"
    for k, v in config_dict.items():
        val = f'"{v}"' if isinstance(v, str) else str(v)
        code += f"    {k}={val},\n"
    code += ")"
    add_code_block(doc, code)

# ── Shared defaults ─────────────────────────────────────────────────
DEFAULTS = dict(batch_size=64, learning_rate=1e-3, device="cpu",
                rank_tol=1e-5, stats_batches=3, warmup_passes=2)
SEEDS = [7, 42, 123]

# ════════════════════════════════════════════════════════════════════
#                        TITLE PAGE
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n').font.size = Pt(24)
r = p.add_run('EXPERIMENT RUNBOOK')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Self-Gated State Space Hybrid Attention\nComplete Run-by-Run Specifications')
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('\nResearcher: Haris\nGenerated: 2026-06-10\nTotal Experiments: 8\nEstimated Total Runs: 168+')
r3.font.size = Pt(11)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#                     TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Global Defaults & Hardware Notes",
    "2. EXP 1 — MQAR Associative Recall Capacity Sweep (72 runs)",
    "3. EXP 2 — Full Ablation Matrix (24 runs)",
    "4. EXP 3 — Scaling Curve: Time & Memory vs Sequence Length (24 runs)",
    "5. EXP 4 — Recurrent vs Parallel Mode Validation (6 runs)",
    "6. EXP 5 — CIFAR-10 Image Classification (12 runs)",
    "7. EXP 6 — Shakespeare Language Modeling (12 runs)",
    "8. EXP 7 — Rank Diagnostics Deep Dive (6 runs)",
    "9. EXP 8 — Multi-Seed Statistical Rigor (10 runs)",
    "10. Master Run Checklist",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#                     GLOBAL DEFAULTS
# ════════════════════════════════════════════════════════════════════
doc.add_heading('1. Global Defaults & Hardware Notes', level=1)
doc.add_paragraph('These defaults are used across ALL experiments unless explicitly overridden in the run specification.')
add_table(doc, ['Parameter', 'Default Value'],
    [['Model dim (D)', '64'], ['Heads (h)', '4'], ['Layers (L)', '2'],
     ['Head dim (d = D/h)', '16'], ['MLP ratio', '2'], ['Window size (w)', '16'],
     ['Learning rate', '1e-3'], ['Batch size', '64'], ['Optimizer', 'AdamW'],
     ['Weight decay', '0.01'], ['Device', 'cpu'], ['SVD rank tolerance', '1e-5'],
     ['Stats batches', '3'], ['Warmup passes', '2'], ['Seeds', '7, 42, 123']])
add_info_box(doc, '⚠️ IMPORTANT', 'Record your hardware specs before starting: CPU model, RAM, GPU (if any), OS version.')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 1 — MQAR CAPACITY SWEEP
# ════════════════════════════════════════════════════════════════════
doc.add_heading('2. EXP 1 — MQAR Associative Recall Capacity Sweep', level=1)
add_info_box(doc, 'Purpose', 'Test memory capacity by sweeping num_pairs and heads. The model must memorize N random key→value pairs and retrieve a specific value when queried.')
add_info_box(doc, 'Hypothesis', 'Memory capacity scales with h × d². As num_pairs exceeds d², accuracy degrades. Increasing heads should restore it.')
add_info_box(doc, 'Expected Figure', 'Heatmap — X: heads (h), Y: num_pairs. Color: validation accuracy. Shows capacity scaling.')

pairs_list = [2, 4, 8, 16, 32, 64]
heads_list = [1, 2, 4, 8]
run_num = 0
total_runs = len(pairs_list) * len(heads_list) * len(SEEDS)
add_info_box(doc, 'Total Runs', f'{len(pairs_list)} pair counts × {len(heads_list)} head counts × {len(SEEDS)} seeds = {total_runs} runs')
doc.add_paragraph()
add_info_box(doc, 'CRITICAL INSTRUCTION', 'Ensure the "Baseline Comparison" checkbox is ENABLED. This will automatically run Softmax, Linear, and RALA alongside the Hybrid model, generating all 4 results simultaneously for each run.')

for np_val in pairs_list:
    for h_val in heads_list:
        for seed in SEEDS:
            run_num += 1
            cfg = dict(task="recall", attention_type="hybrid", num_pairs=np_val,
                       heads=h_val, dim=64, layers=2, recall_vocab=100,
                       sample_limit=10000, epochs=30, seed=seed,
                       **DEFAULTS)
            add_run_entry(doc, run_num, total_runs, cfg,
                          f'pairs={np_val}, heads={h_val}, seed={seed}')
doc.add_paragraph()
add_info_box(doc, 'Record', 'For each run record: val_accuracy, val_loss, memory_rank_ratio, output_rank_ratio, inference_ms')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 2 — ABLATION MATRIX
# ════════════════════════════════════════════════════════════════════
doc.add_heading('3. EXP 2 — Full Ablation Matrix on Associative Recall', level=1)
add_info_box(doc, 'Purpose', 'Isolate the contribution of each architectural component. Full model must beat every ablation.')
add_info_box(doc, 'Hypothesis', 'Full Hybrid > every ablation > vanilla linear. Softmax is a strong baseline but hybrid should match or beat it.')
add_info_box(doc, 'Expected Figure', 'Grouped Bar Chart — X: 8 configurations. Y: validation accuracy. Error bars: ±1 std.')

configs_ablation = [
    ("Full Hybrid",     dict(attention_type="hybrid", use_global=True,  window_size=16, use_output_gate=True,  use_salience_gate=True)),
    ("Global Only",     dict(attention_type="hybrid", use_global=True,  window_size=0,  use_output_gate=True,  use_salience_gate=True)),
    ("Local Only",      dict(attention_type="hybrid", use_global=False, window_size=16, use_output_gate=True,  use_salience_gate=True)),
    ("No φ Gate",       dict(attention_type="hybrid", use_global=True,  window_size=16, use_output_gate=False, use_salience_gate=True)),
    ("No Salience",     dict(attention_type="hybrid", use_global=True,  window_size=16, use_output_gate=True,  use_salience_gate=False)),
    ("Softmax Baseline",dict(attention_type="softmax")),
    ("Linear Baseline", dict(attention_type="linear")),
    ("RALA Baseline",   dict(attention_type="rala")),
]
total_ab = len(configs_ablation) * len(SEEDS)
add_info_box(doc, 'Total Runs', f'{len(configs_ablation)} configs × {len(SEEDS)} seeds = {total_ab} runs')
doc.add_paragraph()

run_num = 0
for name, overrides in configs_ablation:
    for seed in SEEDS:
        run_num += 1
        cfg = dict(task="recall", num_pairs=8, dim=64, heads=4, layers=2,
                   recall_vocab=100, sample_limit=10000, epochs=30, seed=seed,
                   **DEFAULTS)
        cfg.update(overrides)
        add_run_entry(doc, run_num, total_ab, cfg, f'{name}, seed={seed}')
add_info_box(doc, 'Record', 'val_accuracy, val_loss, memory_rank_ratio, output_rank_ratio, inference_ms, param_count')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 3 — SCALING CURVE
# ════════════════════════════════════════════════════════════════════
doc.add_heading('4. EXP 3 — Scaling Curve: Time & Memory vs Sequence Length', level=1)
add_info_box(doc, 'Purpose', 'Measure wall-clock inference time and peak memory as sequence length N grows. Validates O(Nd²) linear complexity claim.')
add_info_box(doc, 'Method', 'Forward pass only (no training). torch.no_grad(). Warmup 3 passes, time 10 passes, average. Report in milliseconds.')
add_info_box(doc, 'Expected Figure', 'Log-Log Line Plot — X: sequence length N. Y: inference time (ms). 4 lines: Hybrid, Softmax, Linear, RALA. Slope=1 is linear, slope=2 is quadratic.')

seq_lengths = [64, 128, 256, 512, 1024, 2048]
scale_models = ["hybrid", "softmax", "linear", "rala"]
total_sc = len(seq_lengths) * len(scale_models)
add_info_box(doc, 'Total Runs', f'{len(seq_lengths)} lengths × {len(scale_models)} models = {total_sc} runs')
doc.add_paragraph()

run_num = 0
for sl in seq_lengths:
    for m in scale_models:
        run_num += 1
        cfg = dict(task="recall", attention_type=m, num_pairs=sl//4, dim=64,
                   heads=4, layers=2, recall_vocab=100, sample_limit=100,
                   epochs=0, seed=7, batch_size=1, learning_rate=1e-3,
                   device="cpu")
        add_run_entry(doc, run_num, total_sc, cfg,
                      f'seq_len≈{sl}, model={m} — INFERENCE ONLY, no training')
add_info_box(doc, 'Record', 'inference_ms (mean of 10 passes), peak_memory_MB (if measurable)')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 4 — RECURRENT vs PARALLEL
# ════════════════════════════════════════════════════════════════════
doc.add_heading('5. EXP 4 — Recurrent vs Parallel Mode Validation', level=1)
add_info_box(doc, 'Purpose', 'Verify recurrent (streaming) mode does not degrade accuracy compared to parallel mode using the same trained weights.')
add_info_box(doc, 'Method', 'Train in parallel mode, then evaluate the SAME checkpoint in both parallel and recurrent mode.')
add_info_box(doc, 'Expected Result', 'Recurrent accuracy within 5% of parallel is acceptable. Beyond 10% is a problem.')

exp4_tasks = [
    ("recall", dict(task="recall", num_pairs=8, recall_vocab=100, sample_limit=10000, seq_len=128)),
    ("shakespeare", dict(task="shakespeare", sample_limit=0, seq_len=128)),
]
exp4_modes = ["parallel", "recurrent"]
total_e4 = len(exp4_tasks) * len(SEEDS)
add_info_box(doc, 'Total Runs', f'{len(exp4_tasks)} tasks × {len(SEEDS)} seeds = {total_e4} training runs, then evaluate each in both modes')
doc.add_paragraph()

run_num = 0
for task_name, task_cfg in exp4_tasks:
    for seed in SEEDS:
        run_num += 1
        cfg = dict(attention_type="hybrid", dim=64, heads=4, layers=2,
                   epochs=30, seed=seed, mode="parallel", **DEFAULTS)
        cfg.update(task_cfg)
        add_run_entry(doc, run_num, total_e4, cfg,
                      f'TRAIN: {task_name}, seed={seed}, mode=parallel')
        p = doc.add_paragraph()
        r = p.add_run(f'    → Then EVALUATE this checkpoint with mode="recurrent" and compare.')
        r.font.size = Pt(8); r.font.italic = True
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
add_info_box(doc, 'Record', 'For EACH mode: val_accuracy, val_loss, memory_rank_ratio, output_rank_ratio, inference_ms. Compute Δ (difference).')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 5 — CIFAR-10
# ════════════════════════════════════════════════════════════════════
doc.add_heading('6. EXP 5 — CIFAR-10 Image Classification', level=1)
add_info_box(doc, 'Purpose', 'Secondary validation on real-world image data. Demonstrates the architecture works on natural data beyond synthetic benchmarks.')
add_info_box(doc, 'Goal', 'Hybrid should match or exceed softmax accuracy while using less compute. Even matching softmax is a win because of linear complexity.')

cifar_models = ["hybrid", "softmax", "linear", "rala"]
total_c = len(cifar_models) * len(SEEDS)
add_info_box(doc, 'Total Runs', f'{len(cifar_models)} models × {len(SEEDS)} seeds = {total_c} runs')
doc.add_paragraph()

run_num = 0
for m in cifar_models:
    for seed in SEEDS:
        run_num += 1
        cfg = dict(task="image", dataset="synthetic", attention_type=m, dim=64,
                   heads=4, layers=4, patch_size=4, sample_limit=5000,
                   epochs=30, seed=seed, **DEFAULTS)
        add_run_entry(doc, run_num, total_c, cfg, f'model={m}, seed={seed}')
add_info_box(doc, 'Record', 'val_accuracy, val_loss, memory_rank_ratio, output_rank_ratio, inference_ms, param_count')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 6 — SHAKESPEARE
# ════════════════════════════════════════════════════════════════════
doc.add_heading('7. EXP 6 — Shakespeare Language Modeling', level=1)
add_info_box(doc, 'Purpose', 'Test on natural language (character-level LM). Validates that the forget gate and local window handle real sequential structure. Metric: perplexity (lower = better).')
add_info_box(doc, 'Goal', 'Hybrid perplexity competitive with softmax. Even slightly higher is acceptable if accompanied by faster inference.')

shk_models = ["hybrid", "softmax", "linear", "rala"]
total_s = len(shk_models) * len(SEEDS)
add_info_box(doc, 'Total Runs', f'{len(shk_models)} models × {len(SEEDS)} seeds = {total_s} runs')
doc.add_paragraph()

run_num = 0
for m in shk_models:
    for seed in SEEDS:
        run_num += 1
        cfg = dict(task="shakespeare", attention_type=m, dim=64, heads=4,
                   layers=2, seq_len=128, sample_limit=0, epochs=30,
                   seed=seed, **DEFAULTS)
        add_run_entry(doc, run_num, total_s, cfg, f'model={m}, seed={seed}')
add_info_box(doc, 'Record', 'perplexity (exp(val_loss)), next_char_accuracy, val_loss, output_rank_ratio, inference_ms')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 7 — RANK DEEP DIVE
# ════════════════════════════════════════════════════════════════════
doc.add_heading('8. EXP 7 — Rank Diagnostics Deep Dive', level=1)
add_info_box(doc, 'Purpose', 'Deep dive into rank metrics to prove the φ gate rank-restoration property. Compare per-layer metrics across Full Hybrid vs No-φ-Gate.')
add_info_box(doc, 'Key Prediction', 'Memory rank ratio < 1.0 (bottleneck). Global output rank ≤ memory rank. But final output rank ≈ 1.0 (φ restores it). Without φ gate, output rank collapses.')
add_info_box(doc, 'Expected Figure', 'Grouped Bar Chart per-layer — 3 bars: Memory Rank Ratio, Global Output Rank Ratio, Final Output Rank Ratio. Line at 1.0.')

rank_configs = [
    ("Full Hybrid",  dict(attention_type="hybrid", use_output_gate=True)),
    ("No φ Gate",    dict(attention_type="hybrid", use_output_gate=False)),
]
total_r = len(rank_configs) * len(SEEDS)
add_info_box(doc, 'Total Runs', f'{len(rank_configs)} configs × {len(SEEDS)} seeds = {total_r} runs')
doc.add_paragraph()

run_num = 0
for name, overrides in rank_configs:
    for seed in SEEDS:
        run_num += 1
        cfg = dict(task="recall", num_pairs=8, dim=64, heads=4, layers=4,
                   recall_vocab=100, sample_limit=10000, epochs=30,
                   seed=seed, use_salience_gate=True, use_global=True,
                   window_size=16, **DEFAULTS)
        cfg.update(overrides)
        add_run_entry(doc, run_num, total_r, cfg, f'{name}, seed={seed}')
add_info_box(doc, 'Record', 'PER-LAYER: memory_rank, memory_rank_ratio, global_output_rank, global_output_ratio, output_rank, output_ratio')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    EXP 8 — MULTI-SEED
# ════════════════════════════════════════════════════════════════════
doc.add_heading('9. EXP 8 — Multi-Seed Statistical Rigor', level=1)
add_info_box(doc, 'Purpose', 'Run key experiments with 5 seeds to compute mean ± std. Required for publishable results.')
add_info_box(doc, 'Standard', '3 seeds minimum for conference papers. 5 seeds for stronger claims.')

seeds_5 = [7, 42, 123, 256, 999]
total_ms = len(seeds_5) * 2  # recall + shakespeare
add_info_box(doc, 'Total Runs', f'{len(seeds_5)} seeds × 2 tasks = {total_ms} runs')
doc.add_paragraph()

run_num = 0
for seed in seeds_5:
    run_num += 1
    cfg = dict(task="recall", attention_type="hybrid", num_pairs=8, dim=64,
               heads=4, layers=2, recall_vocab=100, sample_limit=10000,
               epochs=30, seed=seed, use_output_gate=True,
               use_salience_gate=True, use_global=True, window_size=16,
               **DEFAULTS)
    add_run_entry(doc, run_num, total_ms, cfg, f'Recall, Full Hybrid, seed={seed}')

for seed in seeds_5:
    run_num += 1
    cfg = dict(task="shakespeare", attention_type="hybrid", dim=64, heads=4,
               layers=2, seq_len=128, sample_limit=0, epochs=30, seed=seed,
               use_output_gate=True, use_salience_gate=True, use_global=True,
               window_size=16, **DEFAULTS)
    add_run_entry(doc, run_num, total_ms, cfg, f'Shakespeare, Full Hybrid, seed={seed}')
add_info_box(doc, 'Record', 'val_accuracy, val_loss per seed. Compute mean ± std. Check if std < 0.05 for accuracy.')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#    MASTER CHECKLIST
# ════════════════════════════════════════════════════════════════════
doc.add_heading('10. Master Run Checklist', level=1)
doc.add_paragraph('Use this checklist to track progress. Check off each experiment as you complete it.')
doc.add_paragraph()

checklist = [
    ('EXP 1', 'MQAR Capacity Sweep', '72', '~2-5 min each'),
    ('EXP 2', 'Ablation Matrix', '24', '~2-5 min each'),
    ('EXP 3', 'Scaling Curve', '24', '~1 min each (inference only)'),
    ('EXP 4', 'Recurrent vs Parallel', '6+evals', '~5 min each'),
    ('EXP 5', 'CIFAR-10', '12', '~5-10 min each'),
    ('EXP 6', 'Shakespeare LM', '12', '~3-5 min each'),
    ('EXP 7', 'Rank Deep Dive', '6', '~5 min each'),
    ('EXP 8', 'Multi-Seed Rigor', '10', '~3-5 min each'),
]
add_table(doc, ['Exp', 'Name', 'Runs', 'Est. Time/Run', 'Status'],
    [[c[0], c[1], c[2], c[3], '☐ Not Started'] for c in checklist])

doc.add_paragraph()
add_info_box(doc, 'TOTAL ESTIMATED RUNS', '168+ individual experiment runs')
add_info_box(doc, 'TIP', 'Run experiments in order. EXP 1 and EXP 2 are the most critical for the paper. EXP 3-8 are supporting evidence.')

# ── Save ────────────────────────────────────────────────────────────
import os
from pathlib import Path
outpath = Path(__file__).parent / 'Experiment_Runbook.docx'
doc.save(outpath)
print(f'✅ Document saved to: {outpath}')

